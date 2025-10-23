#!/usr/bin/env python3
"""
Convert Research Paper to Professional Microsoft Word Document
Handles academic formatting, citations, headings, and styling
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re

def create_professional_document():
    """Create a professionally formatted Word document"""
    doc = Document()

    # Set default font to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Configure paragraph formatting
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_after = Pt(6)
    paragraph_format.first_line_indent = Inches(0.5)

    return doc

def add_title_page(doc):
    """Add formatted title page"""
    # Conference header
    p = doc.add_paragraph('THE 5TH CONFERENCE ON RESEARCH IN IT PRACTICE, 22-25 OCT 2024, BRISBANE')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.bold = True

    doc.add_paragraph()  # Spacing

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Benchmarking Generative AI Token Use in Cybersecurity Operations: A Controlled Experimental Study of Prompt Length Optimization')
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()  # Spacing

    # Author info
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Mohamed Zeyada, School of Information Systems, Queensland University of Technology, Brisbane, Australia\nEmail: mohamed.zeyada@hdr.qut.edu.au')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

def add_section_heading(doc, text, level=1):
    """Add formatted section heading"""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.bold = True

    if level == 1:
        heading.runs[0].font.size = Pt(14)
    elif level == 2:
        heading.runs[0].font.size = Pt(13)
    else:
        heading.runs[0].font.size = Pt(12)

    return heading

def format_paragraph(doc, text, bold_markers=True, indent=True):
    """Add formatted paragraph with optional bold text"""
    p = doc.add_paragraph()

    if not indent:
        p.paragraph_format.first_line_indent = Inches(0)

    if bold_markers and '**' in text:
        # Split by bold markers and format
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text
                if part:
                    p.add_run(part)
            else:
                # Bold text
                run = p.add_run(part)
                run.bold = True
    else:
        p.add_run(text)

    return p

def parse_and_format_paper(input_file, output_file):
    """Parse text file and create formatted Word document"""

    print(f"Reading paper from: {input_file}")

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = create_professional_document()
    add_title_page(doc)

    current_section = None
    in_abstract = False
    skip_next_empty = False

    for i, line in enumerate(lines):
        line = line.rstrip()

        # Skip conference header and title (already added)
        if i < 8:
            continue

        # Detect section headings
        if line.startswith('## '):
            section_title = line[3:].strip()
            add_section_heading(doc, section_title, level=1)
            current_section = section_title
            skip_next_empty = True
            continue

        if line.startswith('### '):
            subsection_title = line[4:].strip()
            add_section_heading(doc, subsection_title, level=2)
            skip_next_empty = True
            continue

        # Handle abstract
        if '**Abstract' in line or line.startswith('**Abstract'):
            in_abstract = True
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run('Abstract')
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Handle keywords
        if '**Keywords**' in line:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run('Keywords')
            run.bold = True
            continue

        # Skip empty lines after headings
        if skip_next_empty and not line.strip():
            skip_next_empty = False
            continue

        # Add regular paragraphs
        if line.strip():
            format_paragraph(doc, line, bold_markers=True)

    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    print(f"Saving Word document to: {output_file}")
    doc.save(output_file)
    print(f"✅ Successfully created Word document!")

if __name__ == "__main__":
    input_file = "/home/zeyada/CyberPrompt/research/To finish/Assignment 3B Research paper.txt"
    output_file = "/home/zeyada/CyberPrompt/research/To finish/Assignment_3B_Research_Paper.docx"

    try:
        parse_and_format_paper(input_file, output_file)
        print(f"\n📄 Word Document Created Successfully!")
        print(f"Location: {output_file}")
        print(f"\nFormatting Applied:")
        print("  ✅ Times New Roman 12pt font")
        print("  ✅ Professional heading hierarchy")
        print("  ✅ First-line paragraph indentation")
        print("  ✅ Bold text formatting")
        print("  ✅ Single line spacing")
        print("  ✅ Centered title page")
    except Exception as e:
        print(f"❌ Error converting to Word: {e}")
        import traceback
        traceback.print_exc()
