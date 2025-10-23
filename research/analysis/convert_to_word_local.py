#!/usr/bin/env python3
"""
Research Paper to Microsoft Word Converter
Converts the research paper text file to a professionally formatted Word document
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_page_number(doc):
    """Add page numbers to the document"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def format_title_page(doc, title, authors, conference, email):
    """Create a professional title page"""
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Authors
    authors_para = doc.add_paragraph()
    authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors_para.add_run(authors)
    authors_run.font.size = Pt(14)
    authors_run.font.name = 'Times New Roman'
    
    # Email
    email_para = doc.add_paragraph()
    email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_run = email_para.add_run(email)
    email_run.font.size = Pt(12)
    email_run.font.name = 'Times New Roman'
    
    # Add space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Conference
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run(conference)
    conf_run.font.size = Pt(12)
    conf_run.font.name = 'Times New Roman'
    conf_run.font.bold = True
    
    # Page break
    doc.add_page_break()

def format_abstract(doc, abstract_text):
    """Format the abstract section"""
    # Abstract heading
    abstract_heading = doc.add_heading('Abstract', level=1)
    abstract_heading.style.font.name = 'Times New Roman'
    abstract_heading.style.font.size = Pt(12)
    abstract_heading.style.font.bold = True
    
    # Abstract text
    abstract_para = doc.add_paragraph(abstract_text)
    abstract_para.style.font.name = 'Times New Roman'
    abstract_para.style.font.size = Pt(12)
    abstract_para.paragraph_format.first_line_indent = Inches(0.5)
    abstract_para.paragraph_format.space_after = Pt(6)

def format_keywords(doc, keywords_text):
    """Format keywords section"""
    keywords_para = doc.add_paragraph()
    keywords_para.style.font.name = 'Times New Roman'
    keywords_para.style.font.size = Pt(12)
    
    # Add "Keywords:" label
    keywords_run = keywords_para.add_run("Keywords: ")
    keywords_run.font.bold = True
    
    # Add keywords
    keywords_para.add_run(keywords_text)

def format_section_heading(doc, heading_text, level=1):
    """Format section headings"""
    heading = doc.add_heading(heading_text, level=level)
    heading.style.font.name = 'Times New Roman'
    heading.style.font.size = Pt(12)
    heading.style.font.bold = True
    
    if level == 1:
        heading.style.font.size = Pt(14)
    elif level == 2:
        heading.style.font.size = Pt(13)
    else:
        heading.style.font.size = Pt(12)

def format_paragraph(doc, text):
    """Format regular paragraphs"""
    para = doc.add_paragraph()
    para.style.font.name = 'Times New Roman'
    para.style.font.size = Pt(12)
    para.paragraph_format.first_line_indent = Inches(0.5)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    # Handle bold text
    text = re.sub(r'\*\*(.*?)\*\*', r'<bold>\1</bold>', text)
    
    # Split text and add runs
    parts = re.split(r'<bold>(.*?)</bold>', text)
    for i, part in enumerate(parts):
        if i % 2 == 0:  # Regular text
            if part.strip():
                para.add_run(part)
        else:  # Bold text
            run = para.add_run(part)
            run.font.bold = True
    
    return para

def format_references(doc, references_text):
    """Format references section"""
    format_section_heading(doc, "References", level=1)
    
    # Split references by number
    ref_pattern = r'\[(\d+)\]'
    references = re.split(ref_pattern, references_text)
    
    for i in range(1, len(references), 2):
        if i + 1 < len(references):
            ref_num = references[i]
            ref_text = references[i + 1].strip()
            
            ref_para = doc.add_paragraph()
            ref_para.style.font.name = 'Times New Roman'
            ref_para.style.font.size = Pt(12)
            ref_para.paragraph_format.hanging_indent = Inches(0.5)
            ref_para.paragraph_format.space_after = Pt(3)
            
            # Add reference number
            ref_run = ref_para.add_run(f"[{ref_num}] ")
            ref_run.font.bold = True
            
            # Add reference text
            ref_para.add_run(ref_text)

def convert_research_paper_to_word():
    """Main conversion function"""
    # Read the research paper
    with open('/home/zeyada/CyberPrompt/research/To finish/Assignment 3B Research paper.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Parse content
    lines = content.split('\n')
    
    # Extract title page information
    title = "Benchmarking Generative AI Token Use in Cybersecurity Operations: A Controlled Experimental Study of Prompt Length Optimization"
    authors = "Mohamed Zeyada"
    conference = "THE 5TH CONFERENCE ON RESEARCH IN IT PRACTICE, 22-25 OCT 2024, BRISBANE"
    email = "Email: mohamed.zeyada@hdr.qut.edu.au"
    
    # Create title page
    format_title_page(doc, title, authors, conference, email)
    
    # Process content
    current_section = ""
    current_text = ""
    in_abstract = False
    in_keywords = False
    in_references = False
    references_text = ""
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines and title page content
        if not line or line.startswith("THE 5TH CONFERENCE") or line.startswith("Email:"):
            i += 1
            continue
        
        # Handle abstract
        if line == "Abstract":
            in_abstract = True
            i += 1
            abstract_text = ""
            while i < len(lines) and lines[i].strip() != "Keywords:":
                abstract_text += lines[i] + " "
                i += 1
            format_abstract(doc, abstract_text.strip())
            continue
        
        # Handle keywords
        if line == "Keywords:":
            in_keywords = True
            i += 1
            keywords_text = ""
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("1."):
                keywords_text += lines[i] + " "
                i += 1
            format_keywords(doc, keywords_text.strip())
            continue
        
        # Handle references
        if line == "References":
            in_references = True
            i += 1
            while i < len(lines):
                references_text += lines[i] + "\n"
                i += 1
            format_references(doc, references_text)
            break
        
        # Handle section headings
        if re.match(r'^\d+\.', line) or line.startswith("**") and line.endswith("**"):
            # Save previous section
            if current_text.strip():
                format_paragraph(doc, current_text.strip())
            
            # Process new section
            if re.match(r'^\d+\.', line):
                # Numbered section
                section_text = re.sub(r'^\d+\.\s*', '', line)
                format_section_heading(doc, section_text, level=1)
            elif line.startswith("**") and line.endswith("**"):
                # Bold subsection
                section_text = line[2:-2]
                format_section_heading(doc, section_text, level=2)
            
            current_text = ""
        else:
            # Regular paragraph content
            if line:
                current_text += line + " "
        
        i += 1
    
    # Add any remaining text
    if current_text.strip():
        format_paragraph(doc, current_text.strip())
    
    # Add page numbers
    add_page_number(doc)
    
    # Save document
    output_path = '/home/zeyada/CyberPrompt/research/To finish/Assignment_3B_Research_Paper.docx'
    doc.save(output_path)
    
    print(f"✅ Research paper successfully converted to Word document!")
    print(f"📄 Output file: {output_path}")
    print(f"📊 Document includes:")
    print(f"   - Professional title page")
    print(f"   - Properly formatted abstract and keywords")
    print(f"   - Academic section headings")
    print(f"   - Times New Roman font throughout")
    print(f"   - Proper paragraph indentation")
    print(f"   - Formatted references")
    print(f"   - Page numbers")
    
    return output_path

if __name__ == "__main__":
    try:
        convert_research_paper_to_word()
    except Exception as e:
        print(f"❌ Error converting to Word: {e}")
        print("Make sure python-docx is installed: pip install python-docx")
